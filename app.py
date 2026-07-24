import streamlit as st
import openpyxl
from openpyxl.styles import PatternFill, Alignment
import io
import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import qrcode
from PIL import Image
from streamlit_geolocation import streamlit_geolocation

# Konfigurasi Halaman (Harus diletakkan paling atas)
st.set_page_config(page_title="Aplikasi Survey Lapangan", layout="centered", page_icon="📱")

# Membuat Menu Navigasi di Samping (Sidebar)
st.sidebar.title("📱 Menu Aplikasi")
pilihan_menu = st.sidebar.radio("Pilih Modul Kerja:", ["📝 Pertanyaan Survey (Excel)", "📷 Berita Acara Survey (Word)"])

# ==========================================
# MENU 1: PERTANYAAN SURVEY (EXCEL)
# ==========================================
if pilihan_menu == "📝 Pertanyaan Survey (Excel)":
    st.title("📋 Form Pertanyaan Survey")
    st.markdown("### 👤 Informasi Anggota")
    nama_anggota = st.text_input("Nama dan Nomor Anggota", placeholder="Ketik nama dan nomor anggota di sini...")
    st.markdown("---")

    file_template = "Template_Survey.xlsx" 

    if not os.path.exists(file_template):
        st.error(f"⚠️ File '{file_template}' tidak ditemukan di sistem.")
    else:
        wb = openpyxl.load_workbook(file_template)
        sheet = wb.active
        
        jawaban_user = {}
        sel_warna_input = [] 
        
        for row in range(4, sheet.max_row + 1):
            kategori = sheet.cell(row=row, column=3).value
            pertanyaan = sheet.cell(row=row, column=4).value
            sel_jawaban = sheet.cell(row=row, column=5)
            
            if type(sel_jawaban).__name__ == 'MergedCell':
                continue
                
            nilai_bawaan = sel_jawaban.value
            warna_sel = sel_jawaban.fill.start_color if sel_jawaban.fill else None
            
            if kategori and isinstance(kategori, str) and kategori.strip() != "":
                st.markdown(f"<br>### 📌 {kategori.upper()}", unsafe_allow_html=True)
                
            if pertanyaan and isinstance(pertanyaan, str) and pertanyaan.strip() != "":
                is_yellow = False
                is_green = False
                
                if warna_sel and warna_sel.type == 'theme':
                    if warna_sel.theme == 7:
                        is_yellow = True
                        sel_warna_input.append(sel_jawaban)
                    elif warna_sel.theme == 9:
                        is_green = True
                        sel_warna_input.append(sel_jawaban)
                        
                if is_green and isinstance(nilai_bawaan, str):
                    pilihan_mentah = nilai_bawaan.split('/')
                    pilihan_bersih = [p.strip() for p in pilihan_mentah if p.strip()] 
                    pilihan_bersih.insert(0, "- Pilih Salah Satu -")
                    jawaban = st.selectbox(label=pertanyaan, options=pilihan_bersih, key=f"baris_{row}")
                    if jawaban != "- Pilih Salah Satu -":
                        jawaban_user[row] = jawaban
                        
                elif is_yellow:
                    teks_petunjuk = str(nilai_bawaan).strip() if nilai_bawaan else ""
                    jawaban = st.text_input(label=pertanyaan, placeholder=teks_petunjuk, key=f"baris_{row}")
                    if jawaban:
                        jawaban_user[row] = jawaban
                else:
                    st.markdown(f"**{pertanyaan}**")
                    
        st.markdown("---")
        
        if st.button("💾 Simpan Data & Buat Excel"):
            no_fill = PatternFill(fill_type=None)
            for row, isi_jawaban in jawaban_user.items():
                sel = sheet.cell(row=row, column=5)
                sel.value = isi_jawaban
                
            for sel in sel_warna_input:
                sel.fill = no_fill
                
            cell_b2 = sheet.cell(row=2, column=2)
            nama_bersih = nama_anggota.strip() if nama_anggota else "Tanpa_Nama"
            cell_b2.value = f"Berita Acara Survey : {nama_bersih}"
            cell_b2.fill = no_fill  
            
            sheet.merge_cells(start_row=2, start_column=2, end_row=2, end_column=6)
            cell_b2.alignment = Alignment(horizontal='center', vertical='center')
                        
            for r in range(1, sheet.max_row + 1):
                sheet.row_dimensions[r].height = 30.05
            if sheet.max_row >= 1:
                sheet.row_dimensions[1].height = 13.05
            if sheet.max_row >= 86:
                sheet.row_dimensions[86].height = 13.05
                    
            output = io.BytesIO()
            wb.save(output)
            excel_data = output.getvalue()
            
            nama_file_aman = "".join(c for c in nama_bersih if c.isalnum() or c in (' ', '_', '-')).strip()
            nama_file_final = f"Pertanyaan Survey {nama_file_aman}.xlsx"
            
            st.success(f"✅ Excel Siap! Nama File: **{nama_file_final}**")
            st.download_button(label="📥 Unduh File Excel", data=excel_data, file_name=nama_file_final, mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# ==========================================
# MENU 2: BERITA ACARA SURVEY (WORD DINAMIS)
# ==========================================
elif pilihan_menu == "📷 Berita Acara Survey (Word)":
    st.title("📷 Form Berita Acara & Lampiran Foto")
    st.write("Sistem otomatis penggabungan foto ke format Word multihalaman.")
    
    nama_nasabah = st.text_input("Nama Anggota (Untuk penamaan file)", placeholder="Contoh: Budi Santoso")
    lokasi_teks = st.text_area("Alamat Lengkap (Ditampilkan di Halaman 1)", placeholder="Contoh: Jl. Tembakau 3 No.19 RT.02/01...")
    
    # --- FITUR LOKASI OTOMATIS ---
    st.markdown("---")
    st.write("### 📍 Titik Lokasi / Koordinat")
    st.write("Ketuk tombol **'Get Location'** di bawah ini untuk mengambil titik GPS otomatis. *(Pastikan browser HP Anda diizinkan mengakses lokasi)*")
    
    lokasi_gps = streamlit_geolocation()
    koordinat_otomatis = ""
    
    if lokasi_gps and lokasi_gps.get('latitude') and lokasi_gps.get('longitude'):
        lat = lokasi_gps['latitude']
        lon = lokasi_gps['longitude']
        koordinat_otomatis = f"https://maps.google.com/?q={lat},{lon}"
        st.success("✅ Lokasi berhasil ditemukan!")
        
    koordinat = st.text_input("Hasil Titik Koordinat / Link Maps (Untuk QR Code):", value=koordinat_otomatis, placeholder="Contoh: -6.275, 106.845")
    
    st.markdown("---")
    
    jumlah_halaman = st.number_input("Berapa lembar halaman foto yang dibutuhkan?", min_value=1, max_value=5, value=1, step=1)
    
    data_halaman = []
    
    for i in range(jumlah_halaman):
        st.markdown(f"### 📑 Halaman {i+1}")
        
        judul_default = "FOTO RUMAH YBS" if i == 0 else f"FOTO TAMBAHAN {i+1}"
        judul_halaman = st.text_input(f"Judul Halaman {i+1}", value=judul_default, key=f"judul_{i}")
        
        tab1, tab2 = st.tabs(["📁 Pilih dari Galeri", "📸 Gunakan Kamera"])
        with tab1:
            foto_galeri = st.file_uploader(f"Pilih minimal 2 foto untuk {judul_halaman}", type=['jpg', 'jpeg', 'png'], accept_multiple_files=True, key=f"galeri_{i}")
        with tab2:
            foto_kamera = st.camera_input(f"Kamera untuk {judul_halaman}", key=f"kamera_{i}")
            
        data_halaman.append({
            "judul": judul_halaman,
            "foto_g": foto_galeri,
            "foto_k": foto_kamera
        })
        st.write("") 

    st.markdown("---")
    
    if st.button("📄 Buat Dokumen Word"):
        doc = Document()
        
        # --- PENGATURAN UKURAN KERTAS FOLIO / F4 ---
        # Folio = 8.5 x 13 inches (21.59 cm x 33.02 cm)
        for section in doc.sections:
            section.page_width = Cm(21.59)
            section.page_height = Cm(33.02)
        
        # Pengaturan Font Default (Arial, 11) untuk teks biasa
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        for idx, halaman in enumerate(data_halaman):
            semua_foto = halaman["foto_g"] if halaman["foto_g"] else []
            if halaman["foto_k"]:
                semua_foto.append(halaman["foto_k"])
                
            if idx > 0:
                doc.add_page_break()
                
            # A. Mencetak Judul Halaman (Arial, 14, Bold)
            p_judul = doc.add_paragraph()
            p_judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_judul = p_judul.add_run(halaman["judul"])
            run_judul.bold = True
            run_judul.font.name = 'Arial'
            run_judul.font.size = Pt(14)
            
            # B. Mencetak Foto (6x8 cm, jeda presisi)
            p_foto = doc.add_paragraph()
            p_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if len(semua_foto) == 0:
                p_foto.add_run("*(Tidak ada foto yang dilampirkan pada halaman ini)*")
            else:
                for j, foto in enumerate(semua_foto):
                    img_stream = io.BytesIO(foto.read())
                    run_foto = p_foto.add_run()
                    run_foto.add_picture(img_stream, width=Cm(6), height=Cm(8))
                    
                    posisi = j + 1
                    if posisi % 2 != 0 and posisi < len(semua_foto):
                        run_foto.add_text("    ") # 4 Spasi untuk menyamping
                    elif posisi % 2 == 0 and posisi < len(semua_foto):
                        run_foto.add_text("\n") # 1x Enter untuk baris baru ke bawah (DIUBAH)
            
            # C. Mencetak QR Code & Alamat (HANYA DITAMPILKAN DI HALAMAN 1)
            if idx == 0:
                if koordinat:
                    qr = qrcode.QRCode(box_size=10, border=2)
                    qr.add_data(koordinat)
                    qr.make(fit=True)
                    img_qr = qr.make_image(fill_color="black", back_color="white")
                    
                    qr_stream = io.BytesIO()
                    img_qr.save(qr_stream, format="PNG")
                    qr_stream.seek(0)
                    
                    p_qr = doc.add_paragraph()
                    p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_qr = p_qr.add_run()
                    run_qr.add_picture(qr_stream, width=Cm(5), height=Cm(5))
                    
                if lokasi_teks:
                    # Mencetak Alamat Manual (Arial, 12, Regular)
                    p_alamat = doc.add_paragraph()
                    p_alamat.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_alamat = p_alamat.add_run(lokasi_teks)
                    run_alamat.font.name = 'Arial'
                    run_alamat.font.size = Pt(12)
                    
        # Menyimpan output
        doc_stream = io.BytesIO()
        doc.save(doc_stream)
        doc_data = doc_stream.getvalue()
        
        nama_file_word = f"Berita Acara Survey {nama_nasabah}.docx" if nama_nasabah else "Berita_Acara_Survey.docx"
        
        st.success(f"✅ Dokumen Word (Kertas Folio) Siap! Nama File: **{nama_file_word}**")
        st.download_button(label="📥 Unduh File Word", data=doc_data, file_name=nama_file_word, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")